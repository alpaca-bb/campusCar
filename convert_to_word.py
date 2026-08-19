# -*- coding: utf-8 -*-
"""
将 PROJECT_SUMMARY.md 转换为 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def parse_markdown_to_word(md_file, docx_file):
    """解析 Markdown 文件并生成 Word 文档"""
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    
    # 读取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 标题
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
        
        # 表格
        elif line.startswith('|') and i + 1 < len(lines) and '|---' in lines[i + 1]:
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # 解析表格
            if len(table_lines) >= 2:
                # 跳过分隔行（第二行）
                header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                
                # 创建表格
                table = doc.add_table(rows=1, cols=len(header_cells))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'
                
                # 设置表头
                hdr_cells = table.rows[0].cells
                for j, cell_text in enumerate(header_cells):
                    hdr_cells[j].text = cell_text
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # 添加数据行
                for row_idx in range(2, len(table_lines)):  # 跳过第二行（分隔行）
                    if row_idx >= len(table_lines):
                        break
                    cells = [cell.strip() for cell in table_lines[row_idx].split('|')[1:-1]]
                    if cells:
                        row_cells = table.add_row().cells
                        for j, cell_text in enumerate(cells):
                            if j < len(row_cells):
                                row_cells[j].text = cell_text
                
                # 添加空行
                doc.add_paragraph()
                continue
        
        # 列表项
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
        
        # 有序列表
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
        
        # 代码块
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            # 添加代码段落
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
            # 设置灰色背景
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F5F5F5')
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        
        # 普通段落
        elif line and not line.startswith('---'):
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        
        i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"✅ Word 文档已生成: {docx_file}")

def add_formatted_text(paragraph, text):
    """添加带格式的文本（处理加粗、代码等）"""
    # 处理加粗文本 **text**
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        if idx % 2 == 1:  # 加粗
            run = paragraph.add_run(part)
            run.bold = True
        else:
            # 处理行内代码 `code`
            code_parts = re.split(r'`(.+?)`', part)
            for cidx, code_part in enumerate(code_parts):
                if not code_part:
                    continue
                if cidx % 2 == 1:  # 代码
                    run = paragraph.add_run(code_part)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                else:
                    paragraph.add_run(code_part)

if __name__ == '__main__':
    md_file = r'e:\CyberProject\CampusCar\PROJECT_SUMMARY.md'
    docx_file = r'e:\CyberProject\CampusCar\PROJECT_SUMMARY.docx'
    parse_markdown_to_word(md_file, docx_file)
